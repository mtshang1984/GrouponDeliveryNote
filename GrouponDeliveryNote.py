from ast import If
import os
from pathlib import Path
import re
import time

import pandas as pd #需要使用pandas库
import numpy as np  #需要使用numpy库
import json
from docx import Document  #需要使用python-docx库
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldStart = create_element('w:fldChar')
    create_attribute(fldStart, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'separate')

    fldChar2 = create_element('w:t')
    fldChar2.text = "2"

    fldEnd = create_element('w:fldChar')
    create_attribute(fldEnd, 'w:fldCharType', 'end')

    run._r.append(fldStart)

    run._r.append(instrText)
    run._r.append(fldChar1)
    run._r.append(fldChar2)

    run._r.append(fldEnd)


# 设置word表格列宽度
def set_column_width(table, columns, width_cm):
    col = table.columns[columns]
    for cell in col.cells:
        cell.width = Cm(width_cm)

# 删除word表格列
def delete_column_in_table(table, columns):
    col = table.columns[columns]
    for cell in col.cells:
        cell._element.getparent().remove(cell._element)

# 设置word表格单元格文字内容和格式
def set_cell_text(row, index_row, text, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    row[index_row].text = str(text)
    row[index_row].paragraphs[0].alignment = alignment
    row[index_row].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# 增加一个word表格列出商品派送名细
def add_building_order_table(building_order_data, this_document, groupon_owner, product_name, excel_column_name, building_number, max_building_number, number_line_in_page, max_row_number_per_page, if_hide_phone_number=True,if_upstream_park=False):
    # 本楼栋的产品(按指定顺序排列)，产品种类数

    product_name_for_building = sorted(building_order_data[excel_column_name["product_name"]].unique(
    ), key=lambda x: list(product_name[:, 0]).index(x))
    number_product_name = len(product_name_for_building)

    # 预估本楼栋派送单所需要占用的行数
    number_row_predict_for_building = building_order_data.shape[0] + \
        3*number_product_name+1

    # 如果增加本楼栋的派送单会导致换页，则先换页
    if number_line_in_page+number_row_predict_for_building > max_row_number_per_page:
        this_document.add_page_break()
        number_line_in_page = 0

    for i in range(product_name.shape[0]):
        # 按商品名提取数据
        product_building_order_data = building_order_data.loc[building_order_data[excel_column_name["product_name"]] == product_name[i][0]].sort_values(
            by=[excel_column_name["room_number"]])

        # 如果该商品订单数量不为0，则增加该商品在本楼栋的派送单
        if(product_building_order_data.shape[0] > 0):
            # 预估本楼栋本商品派送单所需要占用的行数
            if product_name[i][0] != product_name_for_building[number_product_name-1]:
                number_row_predict_for_product = product_building_order_data.shape[0]+2 + 1
            else:
                number_row_predict_for_product = product_building_order_data.shape[0]+2 + 3

            # 如果增加表格会导致派送单中间换页，则先接换页
            if number_line_in_page+number_row_predict_for_product > max_row_number_per_page:
                this_document.add_page_break()
                number_line_in_page = 0

            # 增加一个表格，并设置表格格式
            number_row = 0
            table = this_document.add_table(rows=1, cols=7, style='Table Grid')
            table.autofit = True
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells

            # 设置表格抬头
            set_cell_text(hdr_cells, 0, '微信名')
            set_cell_text(hdr_cells, 1, '收货人')
            set_cell_text(hdr_cells, 2, '联系电话')
            set_cell_text(hdr_cells, 3, '楼号')
            set_cell_text(hdr_cells, 4, '房号')
            set_cell_text(hdr_cells, 5, '数量')
            set_cell_text(hdr_cells, 6, groupon_owner+'备注')

            # 设置表格正文内容
            number_of_order = 0
            for index, row in product_building_order_data.iterrows():
                row_cells = table.add_row().cells

                set_cell_text(
                    row_cells, 0, row[excel_column_name["wechat_name"]].encode('utf-8')[0:18].decode('utf-8', errors='ignore'))
                set_cell_text(
                    row_cells, 1, row[excel_column_name["custom_name"]].encode('utf-8')[0:10].decode('utf-8', errors='ignore'))
                set_cell_text(
                    row_cells, 2, row[excel_column_name["phone_number"]])
                
                #增加楼号，如果为555555，则说明订单可能未正确提供楼号，如果为666666，则为商务楼
                if(  row[excel_column_name["building_number"]]==555555):
                    set_cell_text(
                        row_cells, 3, "未能识别的楼号！\n")
                    print("未能识别"+ row[excel_column_name["custom_name"]]+str(row[excel_column_name["phone_number"]])+"订单的楼号，请检查后再试！")
                    exit()
                elif(  row[excel_column_name["building_number"]]==666666):
                    set_cell_text(
                        row_cells, 3, "商务楼")
                else:
                    set_cell_text(
                        row_cells, 3, row[excel_column_name["building_number"]])

                #增加房号，如果为555555，则说明订单可能未正确提供房号，如果为666666，则为别墅区
                if(  row[excel_column_name["room_number"]]==555555):
                    set_cell_text(
                        row_cells, 4, "未能识别的房号！")
                    print("未能识别"+ row[excel_column_name["custom_name"]]+str(row[excel_column_name["phone_number"]])+"订单的房号，请检查后再试！")
                    exit()
                elif(  row[excel_column_name["room_number"]]==666666):
                    set_cell_text(
                        row_cells, 4, "别墅")
                elif(if_upstream_park and  (row[excel_column_name["building_number"]]>=6 and row[excel_column_name["building_number"]]<=36 )):
                    set_cell_text(
                        row_cells, 4, "别墅")
                else:
                    set_cell_text(
                        row_cells, 4, row[excel_column_name["room_number"]])

                #如果为快团团小区团购，则填写数量。如果为普通团购，则不统计数量（普通团购，没有数量这一列，程序暂无法进行解析）
                if excel_column_name["quantity"] in product_building_order_data.columns:
                    set_cell_text(
                        row_cells, 5, row[excel_column_name["quantity"]])
                    number_of_order = number_of_order+int(row[excel_column_name["quantity"] ])
                else:
                    set_cell_text(
                        row_cells, 5, "-")

                #团长备注
                if("remarks" in excel_column_name.keys()):
                    if excel_column_name["remarks"] in product_building_order_data.columns:
                        set_cell_text(
                            row_cells, 6, row[excel_column_name["remarks"]])
                            

            # 最后加一行商品合计行
            row_cells = table.add_row().cells
            # 如果有原订单有数量列，则统计合计数量
            if excel_column_name["quantity"] in product_building_order_data.columns:
                set_cell_text(row_cells, 0, " " +
                            product_name[i][1]+"——合计", WD_ALIGN_PARAGRAPH.LEFT)
            else:
                set_cell_text(row_cells, 0, 
                            product_name[i][1]+"——合计", WD_ALIGN_PARAGRAPH.LEFT)

            if number_of_order==0:
                set_cell_text(row_cells, 5, "-")
            else:
                set_cell_text(row_cells, 5, number_of_order)

            # 设置表格列的宽度
            set_column_width(table, 0, 4)
            set_column_width(table, 3, 2)
            set_column_width(table, 4, 2)
            set_column_width(table, 5, 1.5)
            set_column_width(table, 6, 3)

            # 如果要隐藏手机号，则删除手机号列
            if if_hide_phone_number:
                delete_column_in_table(table, 2)

            # 修改合计行的格式
            for j in range(1, 5):
                if if_hide_phone_number and j == 2:
                    continue
                else:
                    row_cells[0].merge(row_cells[j])

            number_row = number_row + product_building_order_data.shape[0]+2

            # 表格之间增加一行空格
            p = this_document.add_paragraph('')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            number_row = number_row + 1
            number_line_in_page = number_line_in_page+number_row

    # 表格之间增加一行空格
    if(building_number != max_building_number):
        p = this_document.add_paragraph(
            '*****************************************************************************')
        p = this_document.add_paragraph('')
        number_line_in_page = number_line_in_page + 2
    return number_line_in_page

# 输出派送单word文件以供打印


def output_deliverynote_file(data, send_file_name, groupon_owner, product_name, excel_column_name, max_row_number_per_page, page_margin_cm, if_hide_phone_number=True,if_upstream_park=False):
    # 设置字体和段前段后行距
    this_document = Document()
    this_document.styles['Normal'].font.name = u'宋体'
    this_document.styles['Normal']._element.rPr.rFonts.set(
        qn('w:eastAsia'), u'宋体')
    this_document.styles['Normal'].paragraph_format.space_before = Pt(0)
    this_document.styles['Normal'].paragraph_format.space_after = Pt(0)

    # 设置页边距
    sections = this_document.sections
    for section in sections:
        section.top_margin = Cm(page_margin_cm["top_margin"])
        section.bottom_margin = Cm(page_margin_cm["bottom_margin"])
        section.left_margin = Cm(page_margin_cm["left_margin"])
        section.right_margin = Cm(page_margin_cm["right_margin"])

    # 添加页眉
    header = this_document.sections[0].header
    header.is_linked_to_previous = True
    header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header.paragraphs[0].text = "订单确认单（先排楼栋再排品种，本单后面的派送单为先排品种后排楼栋）"

    # 添加页脚页码
    add_page_number(this_document.sections[0].footer.paragraphs[0].add_run())
    this_document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    this_document.sections[0].different_first_page_header_footer = False
    sectPr = this_document.sections[0]._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "1")
    sectPr.append(pgNumType)

    number_line_in_page = 0
    # 按顺序输出每个楼栋的派送单
    # 最大的楼栋号，用于后续程序格式判断用。
    max_building_number = data[excel_column_name["building_number"]].max()
    building_number_list=sorted(data[excel_column_name["building_number"]].unique())
    for i in building_number_list:
    # for i in range(0, max_building_number+1):
    #     if i not in 
        # 获得当前楼栋的订单信息
        building_order_data = data.loc[data[excel_column_name["building_number"]] == i].copy(
        )

        if building_order_data.shape[0] > 0:
            # 增加当前楼栋的表单
            number_line_in_page = add_building_order_table(building_order_data, this_document, groupon_owner, product_name,
                                                           excel_column_name, i, max_building_number, number_line_in_page, max_row_number_per_page, if_hide_phone_number,if_upstream_park)

    p = this_document.add_paragraph(
        '本单据由“小涛”开发的GrouponDeliveryNote程序自动生成，如其他团长有需求，可添加我微信（mtshang1984）免费提供技术支持。')

    # 输出派送单（先排商品，再排楼栋）
    number_blank_product = 0
    if(product_name.shape[0] > 1):

        for j in range(product_name.shape[0]):
            # 获取本套餐订单数据
            product_order_data = data.loc[data[excel_column_name["product_name"]]
                                          == product_name[j][0]].copy()
            if product_order_data.shape[0] == 0:
                number_blank_product = number_blank_product+1
                continue
            # 增加一个小节
            new_section = this_document.add_section(WD_SECTION.ODD_PAGE)
            #设置页码起始为1
            pgNumType.set(ns.qn('w:start'), "1")
            # 设置页眉
            header = this_document.sections[1+j-number_blank_product].header
            header.is_linked_to_previous = False
            header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if excel_column_name["quantity"] in product_order_data.columns:
                header.paragraphs[0].text = "派送单——"+product_name[j][1]+"（共" + \
                str(product_order_data[excel_column_name["quantity"]].sum(
                ))+"份）"
            else:
                header.paragraphs[0].text = "派送单——"+product_name[j][1]


            # 最大的楼栋号，用于后续程序格式判断用。
            max_building_number = data[excel_column_name["building_number"]].max(
            )
            number_line_in_page = 0
            building_number_list=sorted(product_order_data[excel_column_name["building_number"]].unique())
            for i in building_number_list:
                # 获得当前楼栋的订单信息
                building_order_data = product_order_data.loc[product_order_data[excel_column_name["building_number"]] == i].copy(
                )

                if building_order_data.shape[0] > 0:
                    # 增加当前楼栋的表单
                    number_line_in_page = add_building_order_table(building_order_data, this_document, groupon_owner, product_name,
                                                                   excel_column_name, i, max_building_number, number_line_in_page, max_row_number_per_page, if_hide_phone_number,if_upstream_park)

        p = this_document.add_paragraph(
            '本派送单由“小涛”开发的GrouponDeliveryNote程序自动生成，如其他团长有需求，可添加我微信（mtshang1984）免费提供技术支持。')
    else:
        if excel_column_name["quantity"] in data.columns:
            header.paragraphs[0].text = "派送单（共" + \
                str(data[excel_column_name["quantity"]].sum())+"份）"

        else:
            header.paragraphs[0].text = "派送单" 

    this_document.save(send_file_name)

#检查是否为嘉怡水岸小区的团购
def check_if_upstream_park(dataframe,column_name_community,column_name_detail_address):
    if(column_name_community in dataframe.columns ):
        if("嘉怡水岸" in dataframe[column_name_community][0] ):
            return True

    elif(column_name_detail_address in dataframe.columns ):
        if("嘉怡水岸"in dataframe[column_name_detail_address][0] or "紫龙路500" in dataframe[column_name_detail_address][0]) :
            return True
    else:
        return False
#解析详细地址信息至楼号室号
def get_building_number_room_number_from_detial_adrees(detail_address,if_upstream_park=False):
    
    #先以弄对字符串进行分割，正常情况下弄之后的字符串为楼号和室号，将“号，-，-”统一替换为#，以便下一步区分楼号和室号。
    detail_address=detail_address.split("弄")[-1].replace("号","#").replace("-","#").replace("—","#").split("#")

    if len(detail_address)>2:#如果对楼号和室号进行分割后，获得了三个字符串，则只取最后两组字符串作为楼号和室号
        building_number=detail_address[-2]
        room_number=detail_address[-1]
    elif len(detail_address)==2:
        building_number=detail_address[0]
        room_number=detail_address[1]
    elif len(detail_address)==1 and if_upstream_park:
        #如果获得的字符串为1个，则根据字符串内容判断为商务楼或者别墅（此项优化针对嘉怡水岸小区）
        if "商务" in detail_address[0]:
            building_number="商务楼"
            room_number=detail_address[0]
        else:
            building_number=detail_address[0]
            room_number="别墅"
    else:
        print("无法从'"+detail_address+"'解析出楼号和室号，请检查格式")
        exit()


    building_number=re.findall("\d+", (building_number+"-555555").replace(
            "紫龙路500", "嘉怡水岸").replace("商务楼", "666666-"))[0]
    room_number=re.findall("\d+", ("555555-"+ room_number).replace("别墅", "666666-"))[-1]

    return building_number,room_number

# 主程序入口
if __name__ == "__main__":

    start_time = time.time()
    # 检查输入文件，输入文件编码必须为utf-8
    input_file_name = "input.json"
    if os.path.exists(input_file_name):
        program_input = json.load(open(input_file_name, 'r', encoding="utf-8"))
    else:
        print("未找到输入文件："+input_file_name)
        exit()

    # 订单文件
    if "order_file_name" in program_input:
        order_file_name = program_input["order_file_name"]
    else:
        print("未指定订单文件，请确认！")
        exit()

    # 派送单文件
    if "deliverynote_file_name" in program_input:
        deliverynote_file_name = program_input["deliverynote_file_name"]
    else:
        deliverynote_file_name = Path(order_file_name).stem+"派送单.docx"

    # 团长名字
    if "groupon_owner" in program_input:
        groupon_owner = program_input["groupon_owner"]
    else:
        groupon_owner = "团长"

    # 快团团订单表题
    if "excel_column_name" in program_input:
        excel_column_name = program_input["excel_column_name"]
        if "product_name" not in excel_column_name.keys():
            excel_column_name["product_name"] = "商品"
        if "wechat_name" not in excel_column_name.keys():
            excel_column_name["wechat_name"] = "下单人"
        if "custom_name" not in excel_column_name.keys():
            excel_column_name["custom_name"] = "收货人"
        if "phone_number" not in excel_column_name.keys():
            excel_column_name["phone_number"] = "联系电话"
        if "community" not in excel_column_name.keys():
            excel_column_name["community"] = "服务小区"
        if "detail_address" not in excel_column_name.keys():
            excel_column_name["detail_address"] = "详细地址"
        if "building_number" not in excel_column_name.keys():
            excel_column_name["building_number"] = "楼号（如10）"
        if "room_number" not in excel_column_name.keys():
            excel_column_name["room_number"] = "房号（如606）"
        if "quantity" not in excel_column_name.keys():
            excel_column_name["quantity"] = "数量"
        if "remarks" not in excel_column_name.keys():
            excel_column_name["remarks"] = "团长备注"
    else:
        excel_column_name = {
            "product_name": "商品",
            "wechat_name": "下单人",
            "custom_name": "收货人",
            "phone_number": "联系电话",
            "community": "服务小区",
            "detail_address": "详细地址",
            "building_number": "楼号（如10）",
            "room_number": "房号（如606）",
            "quantity": "数量",
            "remarks": "团长备注"
        }
    # 每页最大行数
    if "max_row_number_per_page" in program_input:
        max_row_number_per_page = program_input["max_row_number_per_page"]
    else:
        max_row_number_per_page = 43

    # 页边距
    if "page_margin_cm" in program_input:
        page_margin_cm = program_input["page_margin_cm"]
    else:
        page_margin_cm = {
            "top_margin": 1,
            "bottom_margin": 1,
            "left_margin": 5,
            "right_margin": 1
        }

    # 是否隐藏手机号
    if "if_hide_phone_number" in program_input:
        if_hide_phone_number = program_input["if_hide_phone_number"]
    else:
        if_hide_phone_number = True

    # 读入快团团订单数据
    data = pd.read_excel(order_file_name, keep_default_na=False)

    # 商品列表名称映射表（原名称和派送单中的实际名称的映射关系）
    if "product_name" in program_input:
        product_name = np.array(program_input["product_name"])
    else:
        product_name = np.array([])

    # 如果未提供商品列表名称映射表，则从订单数据重建商品名称列表
    if(product_name.shape[0] == 0):
        product_name_in_order = data[excel_column_name
                                     ["product_name"]].unique()
        number_product_name = len(product_name_in_order)
        
        #如果原订单为小区团购，包含数量列，则在商品名前加数字前缀。
        product_name = [[]]*number_product_name
        for i in range(number_product_name):
            if excel_column_name["quantity"] in data.columns:
                product_name[i] = [product_name_in_order[i],
                                str(i+1)+"-"+product_name_in_order[i]]
            else:
                product_name[i] = [product_name_in_order[i],product_name_in_order[i]]
                max_row_number_per_page = 26

        product_name = np.array(product_name)
    else:
        product_name = np.array(program_input["product_name"])
    #检查是否为嘉怡水岸小区
    if_upstream_park=check_if_upstream_park(data,excel_column_name["community"],excel_column_name["detail_address"])

    #如果未给全楼号和室号，但是有详细地址，则从详细地址信息解析至楼号室号
    if excel_column_name["room_number"] not in data.columns or excel_column_name["building_number"] not in data.columns:
        if excel_column_name["detail_address"]  in data.columns:     
            for index,row in data.iterrows():
                data.at[index,excel_column_name["building_number"]],data.at[index,excel_column_name["room_number"]]=get_building_number_room_number_from_detial_adrees(row[excel_column_name["detail_address"]],if_upstream_park)
            data[excel_column_name["room_number"]] =data[excel_column_name["room_number"]].apply(pd.to_numeric)
            data[excel_column_name["building_number"]] =data[excel_column_name["building_number"]].apply(pd.to_numeric)
        else:
            print("未找到完整地址信息，请确认后再试！")
            exit()
    else:
        # 对房号列进行预处理，过滤掉中文及字符，保留数字部分，别墅用666666作为数字代号，未找到房号的以555555作为数字代号
        if(data[excel_column_name["room_number"]].dtype != np.int32 and data[excel_column_name["room_number"]].dtype != np.int64):
            data[excel_column_name["room_number"]] =("555555-"+ data[excel_column_name["room_number"]].astype(
                str)).str.replace("别墅", "666666-").apply(lambda x: (re.findall("\d+", x)[-1])).apply(pd.to_numeric)#室号转为数字以便排序正确（如果为字符串，1701室会在201室之前，顺序不对）\

                        
        # 对楼号列进行预处理，过滤掉中文及字符，保留数字部分，商务楼用666666作为数字代号，未找到楼号的以555555作为数字代号
        if(data[excel_column_name["building_number"]].dtype != np.int32 and data[excel_column_name["building_number"]].dtype != np.int64):
            data[excel_column_name["building_number"]] = (data[excel_column_name["building_number"]].astype(str)+"-555555").str.replace(
                "500弄", "jiayishuian").str.replace("商务楼", "666666-").apply(lambda x: (re.findall("\d+", x)[0])).apply(pd.to_numeric)#楼号转为数字以便排序正确（如果为字符串，34号会在5号之前，顺序不对）

    #针对嘉怡水岸小区检查房号是否正确
    if if_upstream_park:

            #检查楼号是否正确
            error_data=data[(((data[excel_column_name["building_number"]]>100) &(data[excel_column_name["building_number"]]!=666666)))|(((data[excel_column_name["room_number"]]<101) | ((data[excel_column_name["room_number"]]>1703) &(data[excel_column_name["room_number"]]!=666666)))&((data[excel_column_name["building_number"]]<6)&(data[excel_column_name["building_number"]]>36)))]
            
            for index,row in error_data.iterrows():
                building_number,room_number=(get_building_number_room_number_from_detial_adrees(row[excel_column_name["detail_address"]],if_upstream_park))
                data.at[index,excel_column_name["building_number"]]=int(building_number)
                data.at[index,excel_column_name["room_number"]]=int(room_number)

            
            # if error_data.shape[0]>0:
            #     print(f'{error_data[excel_column_name["custom_name"]].tolist()}房号填写不正确，请检查后再试！')
            #     if_exist_error =True

            # #检查房号是否正确
            # error_data=data[((data[excel_column_name["room_number"]]<101) | ((data[excel_column_name["room_number"]]>1703) &(data[excel_column_name["room_number"]]!=666666)))&((data[excel_column_name["building_number"]]<6)&(data[excel_column_name["building_number"]]>36))]
            # if error_data.shape[0]>0:
            #     print(f'{error_data[excel_column_name["custom_name"]].tolist()}房号填写不正确，请检查后再试！')
            #     if_exist_error =True
                
            # if if_exist_error:
            #     exit()
    print(f"共花费{(time.time()-start_time):0.1f}s完成订单预处理。")


    start_time2=time.time()
    # 输出派送单不带手机号
    output_deliverynote_file(data, deliverynote_file_name, groupon_owner, product_name, excel_column_name,
                             max_row_number_per_page, page_margin_cm, if_hide_phone_number)

    print(f"共花费{(time.time()-start_time2):0.1f}s完成派送单(不带手机号)的生成。")
    
    start_time2=time.time()
    # 输出派送单（带手机号）
    deliverynote_file_name_with_phone_number = Path(deliverynote_file_name).stem \
        + "（带手机号）" \
        + Path(deliverynote_file_name).suffix

    if if_hide_phone_number:
        output_deliverynote_file(data, deliverynote_file_name_with_phone_number, groupon_owner, product_name, excel_column_name, max_row_number_per_page, page_margin_cm, False)
    print(f"共花费{(time.time()-start_time2):0.1f}s完成派送单(带手机号)的生成。")
