import os
from pathlib import Path, PurePosixPath
import re
import time
from tkinter import messagebox

import pandas as pd #需要使用pandas库
import numpy as np  #需要使用numpy库
import json
from docx import Document  
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns

from PyQt5.QtWidgets import QMessageBox,QApplication


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldStart = create_element('w:fldChar')
    create_attribute(fldStart, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'separate')

    fldChar2 = create_element('w:t')
    fldChar2.text = "2"

    fldEnd = create_element('w:fldChar')
    create_attribute(fldEnd, 'w:fldCharType', 'end')

    run._r.append(fldStart)

    run._r.append(instrText)
    run._r.append(fldChar1)
    run._r.append(fldChar2)

    run._r.append(fldEnd)


# 设置word表格列宽度
def set_column_width(table, columns, width_cm):
    col = table.columns[columns]
    for cell in col.cells:
        cell.width = Cm(width_cm)

def set_row_height(table, rows, height_cm):
    row = table.rows[rows]    
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height= Cm(height_cm)

    # for cell in row.cells:
    #     cell.height = Cm(width_cm)

# 删除word表格列
def delete_column_in_table(table, columns):
    col = table.columns[columns]
    for cell in col.cells:
        cell._element.getparent().remove(cell._element)

# 设置word表格单元格文字内容和格式
def set_cell_text(row, index_row, text, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    row[index_row].text = str(text)
    row[index_row].paragraphs[0].alignment = alignment
    row[index_row].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_cell_text_for_lable(row, index_row, text_list, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    row[index_row].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i in range(len(text_list)):
        if i==0:
            para=row[index_row].paragraphs[0]
            para.alignment = alignment
            run = para.add_run(str(text_list[i]))
            run.font.size= Pt(36)
        else:
            para=row[index_row].add_paragraph()
            para.alignment = alignment
            run = para.add_run(str(text_list[i]))
            run.font.size= Pt(11)
            
        run.font.name="微软雅黑"
        
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        run.font.bold=True

# 增加一个word表格列出商品派送名细
def add_building_order_table(building_order_data, this_document, groupon_owner, product_name_list, excel_column_name, building_number, max_building_number, number_line_in_page, max_row_number_per_page, tile_sequence=1,if_hide_phone_number=True,if_upstream_park=False,if_use_pyqt=False,qtwidgets=None):
    # 本楼栋的产品(按指定顺序排列)，产品种类数

    product_name_list_for_building = sorted(building_order_data[excel_column_name["product_name"]].unique(
    ), key=lambda x: list(product_name_list[:, 0]).index(x))
    number_product_name = len(product_name_list_for_building)

    # 预估本楼栋派送单所需要占用的行数
    number_row_predict_for_building = building_order_data.shape[0] + \
        3*number_product_name+1

    # 如果增加本楼栋的派送单会导致换页，则先换页
    if number_line_in_page+number_row_predict_for_building > max_row_number_per_page:
        this_document.add_page_break()
        number_line_in_page = 0

    for i in range(product_name_list.shape[0]):
        # 按商品名提取数据        
        product_building_order_data = building_order_data.loc[building_order_data[excel_column_name["product_name"]] == product_name_list[i][0]]

        # 如果该商品订单数量不为0，则增加该商品在本楼栋的派送单
        if(product_building_order_data.shape[0] > 0):
            #如果超过一个订单，则按房号进行排序
            if(product_building_order_data.shape[0]>1):
                product_building_order_data=product_building_order_data.sort_values(by=[excel_column_name["room_number"]])
            # 预估本楼栋本商品派送单所需要占用的行数
            if product_name_list[i][0] != product_name_list_for_building[number_product_name-1]:
                number_row_predict_for_product = product_building_order_data.shape[0]+2 + 1
            else:
                number_row_predict_for_product = product_building_order_data.shape[0]+2 + 3

            # 如果增加表格会导致派送单中间换页，则先接换页
            if number_line_in_page+number_row_predict_for_product > max_row_number_per_page:
                this_document.add_page_break()
                number_line_in_page = 0

            # 增加一个表格，并设置表格格式
            number_row = 0
            table = this_document.add_table(rows=1, cols=7, style='Table Grid')
            table.autofit = True
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells

            # 设置表格抬头
            if tile_sequence==1:
                index_wechat_name=0
                index_custom_name=1
                index_phone_number=2
                index_building_number=3
                index_room_number=4
                index_quantity=5
                index_remarks=6
            else:
                index_building_number=0
                index_room_number=1
                index_phone_number=2
                index_wechat_name=3
                index_custom_name=4
                index_quantity=5
                index_remarks=6

            set_cell_text(hdr_cells, index_wechat_name, '微信名')
            set_cell_text(hdr_cells, index_custom_name, '收货人')
            set_cell_text(hdr_cells, index_phone_number, '联系电话')
            set_cell_text(hdr_cells, index_building_number, '楼号')
            set_cell_text(hdr_cells, index_room_number, '房号')
            set_cell_text(hdr_cells, index_quantity, '数量')
            set_cell_text(hdr_cells, index_remarks, groupon_owner.encode('utf-8')[0:12].decode('utf-8', errors='ignore')+'备注')

            # 设置表格正文内容
            number_of_order = 0
            for index, row in product_building_order_data.iterrows():
                row_cells = table.add_row().cells

                set_cell_text(
                    row_cells, index_wechat_name, row[excel_column_name["wechat_name"]].encode('utf-8')[0:16].decode('utf-8', errors='ignore'))
                set_cell_text(
                    row_cells, index_custom_name, row[excel_column_name["custom_name"]].encode('utf-8')[0:10].decode('utf-8', errors='ignore'))
                set_cell_text(
                    row_cells, index_phone_number, int(row[excel_column_name["phone_number"]]))
            
                #增加楼号，如果为666666，则为商务楼
                if(row[excel_column_name["building_number"]]==666666):
                    set_cell_text(
                        row_cells, index_building_number, "商务楼")
                else:
                    set_cell_text(
                        row_cells, index_building_number, row[excel_column_name["building_number"]])

                #增加房号，如果为666666，则为别墅区
                if(  row[excel_column_name["room_number"]]==666666):
                    set_cell_text(
                        row_cells, index_room_number, "别墅")
                elif(if_upstream_park and  (row[excel_column_name["building_number"]]>=6 and row[excel_column_name["building_number"]]<=36 )):
                    set_cell_text(
                        row_cells, index_room_number, "别墅")
                else:
                    set_cell_text(
                        row_cells, index_room_number, row[excel_column_name["room_number"]])

                set_cell_text(
                    row_cells, index_quantity, int(row[excel_column_name["quantity"]]))
                number_of_order = number_of_order+int(row[excel_column_name["quantity"] ])

                #团长备注
                if("remarks" in excel_column_name.keys()):
                    if excel_column_name["remarks"] in product_building_order_data.columns:
                        set_cell_text(
                            row_cells,index_remarks, row[excel_column_name["remarks"]])
                            

            # 最后加一行商品合计行
            row_cells = table.add_row().cells
            set_cell_text(row_cells, 0, " " +
                        product_name_list[i][1]+"——合计", WD_ALIGN_PARAGRAPH.LEFT)

            if number_of_order==0:
                set_cell_text(row_cells, index_quantity, "-")
            else:
                set_cell_text(row_cells, index_quantity, number_of_order)

            # 设置表格列的宽度
            set_column_width(table, index_wechat_name, 3.8)
            set_column_width(table, index_building_number, 2)
            set_column_width(table, index_room_number, 2)
            set_column_width(table, index_quantity, 1.5)
            set_column_width(table, index_remarks, 3.2)

            # 如果要隐藏手机号，则删除手机号列
            if if_hide_phone_number:
                delete_column_in_table(table, index_phone_number)

            # 修改合计行的格式
            for j in range(1, index_quantity):
                if if_hide_phone_number and j == index_phone_number:
                    continue
                else:
                    row_cells[0].merge(row_cells[j])

            number_row = number_row + product_building_order_data.shape[0]+2

            # 表格之间增加一行空格
            p = this_document.add_paragraph('')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            number_row = number_row + 1
            number_line_in_page = number_line_in_page+number_row

    # 表格之间增加一行空格
    if(building_number != max_building_number):
        p = this_document.add_paragraph(
            '*****************************************************************************')
        p = this_document.add_paragraph('')
        number_line_in_page = number_line_in_page + 2
    return number_line_in_page

def merge_building_number_and_room_number(building_number,room_number,if_upstream_park,if_minimized_text=False):
    building_number_and_room_number=""
    #增加楼号如果为666666，则为商务楼
    if(  building_number==666666):
        if if_minimized_text:
            building_number_and_room_number="商务"
        else:
            building_number_and_room_number="商务楼"
    else:
        building_number_and_room_number=str(building_number)

    #增加房号，如果为666666，则为别墅区
    if(room_number==666666 or (if_upstream_park and  (building_number>=6 and building_number<=36 ))):

        if if_minimized_text:
            building_number_and_room_number=building_number_and_room_number+"号"
        else:
            building_number_and_room_number=building_number_and_room_number+"号别墅"
    else:
        
        if if_minimized_text:
            building_number_and_room_number=building_number_and_room_number+"-"+str(room_number)
        else:
            building_number_and_room_number=building_number_and_room_number+"号"+str(room_number)+"室"
    return building_number_and_room_number

# 增加一个word表格列出每户商品派送名细
def add_room_order_table(room_order_data, this_document, groupon_owner, product_name_list, excel_column_name,  number_line_in_page, max_row_number_per_page, if_hide_phone_number=True,if_upstream_park=False,if_use_pyqt=False,qtwidgets=None):
    # 预估本户派送单所需要占用的行数
    number_row_predict_for_room = room_order_data.shape[0]+ 2+3

    # 如果增加本户派送单会导致换页，则先换页
    if number_line_in_page+number_row_predict_for_room > max_row_number_per_page:
        this_document.add_page_break()
        number_line_in_page = 0

    # 如果该户商品订单数量不为0，则增加该户的派送单
    if(room_order_data.shape[0] > 0):
        # 按自定义商品顺序排序
        if(room_order_data.shape[0] > 1):
            product_name_dict={}
            for i in range(product_name_list.shape[0]):
                product_name_dict[product_name_list[i,0]]=i

            room_order_data = room_order_data.sort_values(
                by=excel_column_name["product_name"], key=lambda x:x.map(product_name_dict))

        number_row_predict_for_room = room_order_data.shape[0]+2 + 3

        # 如果增加表格会导致派送单中间换页，则先接换页
        if number_line_in_page+number_row_predict_for_room > max_row_number_per_page:
            this_document.add_page_break()
            number_line_in_page = 0

        # 增加一个表格，并设置表格格式
        number_row = 0
        table = this_document.add_table(rows=1, cols=6, style='Table Grid')
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells

        # 设置表格抬头
        index_wechat_name=0
        index_custom_name=1
        index_phone_number=2
        index_product_name=3
        index_quantity=4
        index_remarks=5

        set_cell_text(hdr_cells, index_wechat_name, '微信名')
        set_cell_text(hdr_cells, index_custom_name, '收货人')
        set_cell_text(hdr_cells, index_phone_number, '联系电话')
        set_cell_text(hdr_cells, index_product_name, '商品名称')
        set_cell_text(hdr_cells, index_quantity, '数量')
        set_cell_text(hdr_cells, index_remarks, groupon_owner.encode('utf-8')[0:12].decode('utf-8', errors='ignore')+'备注')

        # 设置表格正文内容
        number_of_order = 0
        for index, row in room_order_data.iterrows():
            row_cells = table.add_row().cells

            wechat_name=row[excel_column_name["wechat_name"]].encode('utf-8')[0:10].decode('utf-8', errors='ignore')
            custom_name=row[excel_column_name["custom_name"]].encode('utf-8')[0:10].decode('utf-8', errors='ignore')
            phone_number=row[excel_column_name["phone_number"]]
            product_name=product_name_list[list(product_name_list[:, 0]).index(row[excel_column_name["product_name"]])][1].encode('utf-8')[0:50].decode('utf-8', errors='ignore')
            
            building_number=row[excel_column_name["building_number"]]
            room_number=row[excel_column_name["room_number"]]

            set_cell_text(
                row_cells, index_wechat_name, wechat_name.encode('utf-8')[0:9].decode('utf-8', errors='ignore'))
            set_cell_text(
                row_cells, index_custom_name, custom_name.encode('utf-8')[0:9].decode('utf-8', errors='ignore'))
            set_cell_text(
                row_cells, index_phone_number, int(phone_number))
            set_cell_text(
                row_cells, index_product_name, product_name)

            #如果为快团团小区团购，则填写数量。如果为普通团购，则不统计数量（普通团购，没有数量这一列，程序暂无法进行解析）
            if excel_column_name["quantity"] in room_order_data.columns:
                set_cell_text(
                    row_cells, index_quantity, int(row[excel_column_name["quantity"]]))
                number_of_order = number_of_order+int(row[excel_column_name["quantity"] ])
            else:
                set_cell_text(
                    row_cells, index_quantity, "-")

            #团长备注
            if("remarks" in excel_column_name.keys()):
                if excel_column_name["remarks"] in room_order_data.columns:
                    set_cell_text(
                        row_cells, index_remarks, row[excel_column_name["remarks"]])
                        
        building_number_and_room_number=merge_building_number_and_room_number(building_number,room_number,if_upstream_park)

        # 最后加一行商品合计行
        row_cells = table.add_row().cells
        set_cell_text(row_cells, 0, f"{building_number_and_room_number}——合计")

        if number_of_order==0:
            set_cell_text(row_cells, index_quantity, "-")
        else:
            set_cell_text(row_cells, index_quantity, number_of_order)

        # 设置表格列的宽度
        set_column_width(table, index_wechat_name, 2)
        set_column_width(table, index_custom_name, 2)
        set_column_width(table, index_product_name, 8)
        set_column_width(table, index_quantity, 1.5)
        set_column_width(table, index_remarks, 3.2)

        # 如果要隐藏手机号，则删除手机号列
        if if_hide_phone_number:
            delete_column_in_table(table, index_phone_number)

        # 修改合计行的格式
        for j in range(1, index_quantity):
            if if_hide_phone_number and j == index_phone_number:
                continue
            else:
                row_cells[0].merge(row_cells[j])

        number_row = number_row + room_order_data.shape[0]+2

        # 表格之间增加一行空格
        p = this_document.add_paragraph('')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p = this_document.add_paragraph(
            '*****************************************************************************')
        p = this_document.add_paragraph('')
        number_row = number_row + 3
        number_line_in_page = number_line_in_page+number_row

    return number_line_in_page

#设置文档第一页格式
def set_first_page(this_document,page_margin_cm,show_sequence):
    this_document.styles['Normal'].font.name = u'宋体'
    this_document.styles['Normal']._element.rPr.rFonts.set(
        qn('w:eastAsia'), u'宋体')
    this_document.styles['Normal'].paragraph_format.space_before = Pt(0)
    this_document.styles['Normal'].paragraph_format.space_after = Pt(0)

    # 设置页边距
    sections = this_document.sections
    for section in sections:
        section.top_margin = Cm(page_margin_cm["top_margin"])
        section.bottom_margin = Cm(page_margin_cm["bottom_margin"])
        section.left_margin = Cm(page_margin_cm["left_margin"])
        section.right_margin = Cm(page_margin_cm["right_margin"])

    if show_sequence!=4:
        add_page_number(this_document.sections[0].footer.paragraphs[0].add_run())    
        this_document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        this_document.sections[0].different_first_page_header_footer = False
        sectPr = this_document.sections[0]._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(ns.qn('w:start'), "1")
        sectPr.append(pgNumType)

# 添加页眉
def set_header(header,is_linked_to_previous,text):    
    header.is_linked_to_previous = is_linked_to_previous
    header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header.paragraphs[0].text = text
    
# 添加页脚页码起始数字
def set_page_number(start_page_number):
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), start_page_number)

def save_document(this_document,filename,if_use_pyqt,qtwidgets):
    try:
        this_document.save(filename)
        return True
    except Exception as e:
        messagebox_text= f"请先关闭文件{filename}，然后点击确定重试！"
        if if_use_pyqt:
            reply=QMessageBox.critical(qtwidgets,"错误对话框",messagebox_text,QMessageBox.Ok | QMessageBox.Cancel)
            if(reply==QMessageBox.Ok):
                return save_document(this_document,filename,if_use_pyqt,qtwidgets)
            else:
                return False
        else:
            print(messagebox_text)
            return False

#增加包含标签的表格
def add_lable_table(this_document):
    
    table = this_document.add_table(rows=5, cols=3, style='Table Grid')
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(2):
        set_column_width(table,i,6.9)
    for i in range(5):
        set_row_height(table,i,5.3)
    return table

# 输出派送单word文件以供打印
def output_deliverynote_file(data, delivery_note_file_name, groupon_owner, product_name_list, excel_column_name, max_row_number_per_page, page_margin_cm,show_sequence=1,tile_sequence=1, if_hide_phone_number=True,if_upstream_park=False,if_use_pyqt=False,qtwidgets=None):
    # 设置字体和段前段后行距
    this_document = Document()
    
    
    # 最大的楼栋号，用于后续程序格式判断用。
    max_building_number = data[excel_column_name["building_number"]].max()

    if(show_sequence==1):
        number_blank_product=0
        for j in range(product_name_list.shape[0]):
            # 获取本套餐订单数据
            product_order_data = data.loc[data[excel_column_name["product_name"]]
                                          == product_name_list[j][0]].copy()
            if product_order_data.shape[0] == 0:
                number_blank_product = number_blank_product+1
                continue
            # 增加一个小节
            if j==0:
                set_first_page(this_document,page_margin_cm,show_sequence)
            else:
                new_section = this_document.add_section(WD_SECTION.ODD_PAGE)
                #设置页码起始为1
                set_page_number("1")

            # 设置页眉
            header = this_document.sections[j-number_blank_product].header

            if excel_column_name["quantity"] in product_order_data.columns:                
                set_header(header,False,"派送单——"+product_name_list[j][1]+"（共" + \
                str(int(product_order_data[excel_column_name["quantity"]].sum(
                )))+"份）")
            else:
                header.paragraphs[0].text = set_header(header,False,"派送单——"+product_name_list[j][1])

            number_line_in_page = 0
            building_number_list=sorted(product_order_data[excel_column_name["building_number"]].unique())
            for i in building_number_list:
                # 获得当前楼栋的订单信息
                product_building_order_data = product_order_data.loc[product_order_data[excel_column_name["building_number"]] == i].copy(
                )

                if product_building_order_data.shape[0] > 0:
                    # 增加当前楼栋的表单
                    number_line_in_page = add_building_order_table(product_building_order_data, this_document, groupon_owner, product_name_list,
                                                                   excel_column_name, i, max_building_number, number_line_in_page, max_row_number_per_page, tile_sequence,if_hide_phone_number,if_upstream_park,if_use_pyqt,qtwidgets)
                    if number_line_in_page ==-1:
                        return False

        p = this_document.add_paragraph(
            '本派送单由“小涛”开发的GrouponDeliveryNote程序自动生成，如其他团长有需求，可添加我微信（mtshang1984）免费提供技术支持。')
    elif(show_sequence==2):
        set_first_page(this_document,page_margin_cm,show_sequence)
        set_header( this_document.sections[0].header,True, "派送单")
        set_page_number("1")

        number_line_in_page = 0
        # 按顺序楼栋-商品-楼号的顺序输出派送单
        building_number_list=sorted(data[excel_column_name["building_number"]].unique())
        for i in building_number_list:
            # 获得当前楼栋的订单信息
            product_building_order_data = data.loc[data[excel_column_name["building_number"]] == i].copy(
            )

            if product_building_order_data.shape[0] > 0:
                # 增加当前楼栋的表单
                number_line_in_page = add_building_order_table(product_building_order_data, this_document, groupon_owner, product_name_list,
                                                            excel_column_name, i, max_building_number, number_line_in_page, max_row_number_per_page, tile_sequence,if_hide_phone_number,if_upstream_park,if_use_pyqt,qtwidgets)
                if number_line_in_page ==-1:
                    return False
        p = this_document.add_paragraph(
            '本单据由“小涛”开发的GrouponDeliveryNote程序自动生成，如其他团长有需求，可添加我微信（mtshang1984）免费提供技术支持。')
    elif(show_sequence==3):
        set_first_page(this_document,page_margin_cm,show_sequence)
        set_header( this_document.sections[0].header,True, "派送单")
        set_page_number("1")

        number_line_in_page = 0
        # 按顺序楼栋-楼号-商品的顺序输出派送单
        building_number_list=sorted(data[excel_column_name["building_number"]].unique())
        for i in building_number_list:
            # 获得当前楼栋的订单信息
            product_building_order_data = data.loc[data[excel_column_name["building_number"]] == i].copy(
            )
            if product_building_order_data.shape[0] > 0:
                room_number_list=sorted(product_building_order_data[excel_column_name["room_number"]].unique())
                for j in room_number_list:
                    # 获得当前户的订单信息
                    room_order_data = product_building_order_data.loc[data[excel_column_name["room_number"]] == j].copy()
                    if(room_order_data.shape[0] > 0):
                        number_line_in_page = add_room_order_table(room_order_data, this_document, groupon_owner, product_name_list,excel_column_name, number_line_in_page, max_row_number_per_page, if_hide_phone_number,if_upstream_park,if_use_pyqt,qtwidgets)
                        if number_line_in_page ==-1:
                            return False
        p = this_document.add_paragraph(
            '本单据由“小涛”开发的GrouponDeliveryNote程序自动生成，如其他团长有需求，可添加我微信（mtshang1984）免费提供技术支持。')
    elif(show_sequence==4):
        set_first_page(this_document,page_margin_cm,show_sequence)
        number_label_in_page = 0
        table = add_lable_table(this_document)

        for i in range(product_name_list.shape[0]):
            # 获取本套餐订单数据
            index_in_product=0
            product_order_data = data.loc[data[excel_column_name["product_name"]]== product_name_list[i][0]].sort_values(by = [excel_column_name["building_number"],excel_column_name["room_number"]])
            product_quantity_sum=product_order_data[excel_column_name["quantity"]].sum() 
            groupon_owner_string="-"+groupon_owner.encode('utf-8')[0:12].decode('utf-8', errors='ignore')
            building_number_list=sorted(product_order_data[excel_column_name["building_number"]].unique())
            
            for j in building_number_list:
                # 获得当前楼栋的订单信息
                product_building_order_data = product_order_data.loc[product_order_data[excel_column_name["building_number"]] == j].copy(
                )                
                product_building_quantity_sum=product_building_order_data[excel_column_name["quantity"]].sum()
                
                index_in_product_building=0
                for index, row in product_building_order_data.iterrows():
                    quantity=int(row[excel_column_name["quantity"]])
                    for k in range(quantity):
                        if number_label_in_page>=15:
                            number_label_in_page=number_label_in_page-15
                            table = add_lable_table(this_document)

                        index_row=int((number_label_in_page)/3)
                        index_column=number_label_in_page+1-3*(index_row)-1
                        row_cells = table.rows[index_row].cells

                        custom_name=row[excel_column_name["custom_name"]].encode('utf-8')[0:10].decode('utf-8', errors='ignore')
                        phone_number=row[excel_column_name["phone_number"]]
                        product_name=product_name_list[list(product_name_list[:, 0]).index(row[excel_column_name["product_name"]])][1].encode('utf-8')[0:50].decode('utf-8', errors='ignore')
                        
                        building_number_and_room_number=merge_building_number_and_room_number(row[excel_column_name["building_number"]],row[excel_column_name["room_number"]],if_upstream_park,True)
                        
                        if if_hide_phone_number:
                            additional_string=""
                        else:
                            additional_string=str(phone_number)
                        set_cell_text_for_lable(row_cells, index_column, [building_number_and_room_number,custom_name+additional_string,product_name,f'总{index_in_product+1}/{product_quantity_sum}(本楼{index_in_product_building+1}/{product_building_quantity_sum}){groupon_owner_string}'])
                        index_in_product=index_in_product+1
                        index_in_product_building=index_in_product_building+1
                        number_label_in_page=number_label_in_page+1

            #不同品种之间空一格
            number_label_in_page=number_label_in_page+1
    else:
        messagebox_text= f"不支持show_sequence值为{show_sequence}的排序方式！"
        if if_use_pyqt:
            QMessageBox.critical(qtwidgets,"错误对话框",messagebox_text)
        else:
            print(messagebox_text)
        return False

    return save_document(this_document,delivery_note_file_name,if_use_pyqt,qtwidgets)

#检查是否为嘉怡水岸小区的团购
def check_if_upstream_park(dataframe,column_name_community,column_name_detail_address):
    if(column_name_community in dataframe.columns ):
        if("嘉怡水岸" in dataframe[column_name_community][0] ):
            return True

    elif(column_name_detail_address in dataframe.columns ):
        if("嘉怡水岸"in dataframe[column_name_detail_address][0] or "紫龙路500" in dataframe[column_name_detail_address][0] or "龙吴路5899" in dataframe[column_name_detail_address][0]) :
            return True
    else:
        return False
#解析详细地址信息至楼号室号
def get_building_number_room_number_from_detial_adrees(detail_address,if_upstream_park=False,if_use_pyqt=False,qtwidgets=None):
    
    #先以弄对字符串进行分割，正常情况下弄之后的字符串为楼号和室号，将“号，-，-”统一替换为#，以便下一步区分楼号和室号。
    detail_address_list=detail_address.split("弄")[-1].replace("号","#").replace("栋","#").replace("幢","#").replace("-","#").replace("—","#").split("#")

    if len(detail_address_list)>2:#如果对楼号和室号进行分割后，获得了三个字符串，则只取最后两组字符串作为楼号和室号
        building_number=detail_address_list[-2]
        room_number=detail_address_list[-1]
    elif len(detail_address_list)==2:
        if if_upstream_park and "龙吴路5899" in detail_address_list[0]:
            building_number="666666"
            room_number=detail_address_list[1]
        else:
            building_number=detail_address_list[0]
            room_number=detail_address_list[1]
    elif len(detail_address_list)==1 and if_upstream_park:
        #如果获得的字符串为1个，则根据字符串内容判断为商务楼或者别墅（此项优化针对嘉怡水岸小区）
        if "商务" in detail_address_list[0] or "龙吴路5899" in  detail_address_list[0]:
            building_number="商务楼"
            room_number=detail_address_list[0]
        else:
            building_number=detail_address_list[0]
            room_number="别墅"
    else:
        messagebox_text= f'无法从"{detail_address_list}"解析出楼号和室号，请检查格式'
        if if_use_pyqt:
            QMessageBox.critical(qtwidgets,"错误对话框",messagebox_text)
            app = QApplication.instance()
            app.quit()
        else:
            print(messagebox_text)
            exit()


    building_number=re.findall("\d+", (building_number+"-555555").replace(
            "紫龙路500", "嘉怡水岸").replace("商务楼", "666666-"))[0]
    room_number=re.findall("\d+", ("555555-"+ room_number).replace("别墅", "666666-"))[-1]

    return building_number,room_number

#根据参数自动生成派送单文件名称
def generate_deliverynote_file_name(order_file_name,if_hide_phone_number,show_sequence):
        keyword=""
        if(if_hide_phone_number==False):
            keyword="（含手机号）"
        if(show_sequence==1):
            keyword=keyword+"（按商品-楼号-房号排序）"
        elif(show_sequence==2):
            keyword=keyword+"（按楼号-商品-房号排序）"
        elif(show_sequence==3):
            keyword=keyword+"（按楼号-房号-商品排序）"
        elif(show_sequence==4):
            keyword=keyword+"（打印标签）"

        return str(PurePosixPath(order_file_name).parent)+"/"+Path(order_file_name).stem+"派送单"+keyword+".docx"

#主处理程序
def main_program(input_file_name,if_use_pyqt=False,qtwidgets=None):

    start_time = time.time()
    # 检查输入文件，输入文件编码必须为utf-8
    if os.path.exists(input_file_name):
        program_input = json.load(open(input_file_name, 'r', encoding="utf-8"))
    else:
        messagebox_text=f"未找到输入文件：{input_file_name}"
        if if_use_pyqt:
            QMessageBox.critical(qtwidgets,messagebox_text)
            app = QApplication.instance()
            app.quit()
        else:
            print(messagebox_text)
            exit()

    # 团长名字
    if "groupon_owner" in program_input:
        groupon_owner = program_input["groupon_owner"]
    else:
        groupon_owner = "团长"

    # 订单文件
    if "order_file_name" in program_input:
        order_file_name = program_input["order_file_name"]
    else:
        messagebox_text= f"未指定订单文件{order_file_name}，请确认！"
        if if_use_pyqt:
            QMessageBox.critical(qtwidgets,messagebox_text)
            app = QApplication.instance()
            app.quit()
        else:
            print(messagebox_text)            
            exit()


    # 派送单排序方法
    if "show_sequence" in program_input:
        show_sequence = program_input["show_sequence"]
    else:
        show_sequence = 1

    # 表题顺序
    if "title_sequence" in program_input:
        title_sequence = program_input["title_sequence"]
    else:
        title_sequence = 1

    # 快团团订单表题
    if "excel_column_name" in program_input:
        excel_column_name = program_input["excel_column_name"]
        if "product_name" not in excel_column_name.keys():
            excel_column_name["product_name"] = "商品"
        if "wechat_name" not in excel_column_name.keys():
            excel_column_name["wechat_name"] = "下单人"
        if "custom_name" not in excel_column_name.keys():
            excel_column_name["custom_name"] = "收货人"
        if "phone_number" not in excel_column_name.keys():
            excel_column_name["phone_number"] = "联系电话"
        if "community" not in excel_column_name.keys():
            excel_column_name["community"] = "服务小区"
        if "detail_address" not in excel_column_name.keys():
            excel_column_name["detail_address"] = "详细地址"
        if "building_number" not in excel_column_name.keys():
            excel_column_name["building_number"] = "楼号（如10）"
        if "room_number" not in excel_column_name.keys():
            excel_column_name["room_number"] = "房号（如606）"
        if "quantity" not in excel_column_name.keys():
            excel_column_name["quantity"] = "数量"
        if "remarks" not in excel_column_name.keys():
            excel_column_name["remarks"] = "团长备注"
    else:
        excel_column_name = {
            "product_name": "商品",
            "wechat_name": "下单人",
            "custom_name": "收货人",
            "phone_number": "联系电话",
            "community": "服务小区",
            "detail_address": "详细地址",
            "building_number": "楼号（如10）",
            "room_number": "房号（如606）",
            "quantity": "数量",
            "remarks": "团长备注"
        }
    # 每页最大行数
    if "max_row_number_per_page" in program_input:
        max_row_number_per_page = program_input["max_row_number_per_page"]
    else:
        max_row_number_per_page = 43

    # 页边距
    if "page_margin_cm" in program_input:
        page_margin_cm = program_input["page_margin_cm"]
    else:
        if show_sequence ==3:
            page_margin_cm = {
                "top_margin": 1,
                "bottom_margin": 1,
                "left_margin": 3,
                "right_margin": 1
            }
        elif show_sequence ==4:
            page_margin_cm = {
                "top_margin": 0.5,
                "bottom_margin": 0.2,
                "left_margin": 0.5,
                "right_margin": 0.5
            }
        else:
            page_margin_cm = {
                "top_margin": 1,
                "bottom_margin": 1,
                "left_margin": 5,
                "right_margin": 1
            }

    # 是否隐藏手机号
    if "if_hide_phone_number" in program_input:
        if_hide_phone_number = program_input["if_hide_phone_number"]
    else:
        if_hide_phone_number = True

    # 派送单文件
    if "deliverynote_file_name" in program_input:
        deliverynote_file_name = program_input["deliverynote_file_name"]
    else:
        deliverynote_file_name=generate_deliverynote_file_name(order_file_name,if_hide_phone_number,show_sequence)

    # 读入快团团订单数据
    data = pd.read_excel(order_file_name, keep_default_na=False)

    # 判断订单是否快团团社区团购
    if excel_column_name["quantity"] in data.columns:
        if_community_groupon=True
    else:
        if_community_groupon=False
    # 如果不是社区团购，则将普通团购数据格式转换为社区团购表单形式
    if(if_community_groupon==False):
        i=0
        old_data=data.copy()
        old_data[excel_column_name["quantity"]]=1

        data=pd.DataFrame()        
        for index,row in old_data.iterrows():
            product_order_list=row[excel_column_name["product_name"]].split(";\n")
            for product_order in product_order_list:
                spliter_index=product_order.rindex("+")
                product_name=product_order[0:spliter_index]
                quantity=product_order[spliter_index+1:]
                row[excel_column_name["product_name"]]=product_name
                row[excel_column_name["quantity"]]=int(quantity)
                data=data.append(row,ignore_index=True)
    # 商品列表名称映射表（原名称和派送单中的实际名称的映射关系）
    if "product_name_list" in program_input:
        product_name_list = np.array(program_input["product_name_list"])
    else:
        product_name_list = np.array([])

    # 如果未提供商品列表名称映射表，则从订单数据重建商品名称列表
    data[excel_column_name["product_name"]]=data[excel_column_name["product_name"]].str.replace("\n","")
    if(product_name_list.shape[0] == 0):
        product_name_list_in_order = data[excel_column_name
                                     ["product_name"]].unique()
        number_product_name = len(product_name_list_in_order)
        
        #如果原订单为小区团购，包含数量列，则在商品名前加数字前缀。
        product_name_list = [[]]*number_product_name
        for i in range(number_product_name):
            if excel_column_name["quantity"] in data.columns:
                product_name_list[i] = [product_name_list_in_order[i],
                                str(i+1)+"-"+product_name_list_in_order[i]]
            else:
                product_name_list[i] = [product_name_list_in_order[i],product_name_list_in_order[i]]
                max_row_number_per_page = 26

        product_name_list = np.array(product_name_list)
    else:
        product_name_list = np.array(program_input["product_name_list"])
    #检查是否为嘉怡水岸小区
    if_upstream_park=check_if_upstream_park(data,excel_column_name["community"],excel_column_name["detail_address"])

    #如果未给全楼号和室号，但是有详细地址，则从详细地址信息解析至楼号室号
    if excel_column_name["room_number"] not in data.columns or excel_column_name["building_number"] not in data.columns:
        if excel_column_name["detail_address"]  in data.columns:     
            for index,row in data.iterrows():
                data.at[index,excel_column_name["building_number"]],data.at[index,excel_column_name["room_number"]]=get_building_number_room_number_from_detial_adrees(row[excel_column_name["detail_address"]],if_upstream_park,if_use_pyqt,qtwidgets)
            data[excel_column_name["room_number"]] =data[excel_column_name["room_number"]].apply(pd.to_numeric)
            data[excel_column_name["building_number"]] =data[excel_column_name["building_number"]].apply(pd.to_numeric)
        else:
            messagebox_text= "订单中未找到完整地址信息，请确认后再试！"
            if if_use_pyqt:
                QMessageBox.critical(qtwidgets,messagebox_text)
                app = QApplication.instance()
                app.quit()
            else:
                print(messagebox_text)            
                exit()
    else:
        # 对房号列进行预处理，过滤掉中文及字符，保留数字部分，别墅用666666作为数字代号，未找到房号的以555555作为数字代号
        if(data[excel_column_name["room_number"]].dtype != np.int32 and data[excel_column_name["room_number"]].dtype != np.int64):
            data[excel_column_name["room_number"]] =("555555-"+ data[excel_column_name["room_number"]].astype(
                str)).str.replace(".*别墅.*", "-666666",regex=True).apply(lambda x: (re.findall("\d+", x)[-1])).apply(pd.to_numeric)#室号转为数字以便排序正确（如果为字符串，1701室会在201室之前，顺序不对）\

                        
        # 对楼号列进行预处理，过滤掉中文及字符，保留数字部分，商务楼用666666作为数字代号，未找到楼号的以555555作为数字代号
        if(data[excel_column_name["building_number"]].dtype != np.int32 and data[excel_column_name["building_number"]].dtype != np.int64):
            data[excel_column_name["building_number"]] = (data[excel_column_name["building_number"]].astype(str)+"-555555").str.replace(
                "紫龙500", "嘉怡水岸").str.replace(".*龙吴路5899.*", "666666-",regex=True).str.replace(".*商务楼.*", "666666-",regex=True).apply(lambda x: (re.findall("\d+", x)[0])).apply(pd.to_numeric)#楼号转为数字以便排序正确（如果为字符串，34号会在5号之前，顺序不对）

    #针对嘉怡水岸小区检查房号是否正确
    if if_upstream_park:

        #检查楼号是否正确，如果不正确，则从地址中重新获取
        error_data=data[(((data[excel_column_name["building_number"]]>100) &(data[excel_column_name["building_number"]]!=666666)))|(((data[excel_column_name["room_number"]]<101) | ((data[excel_column_name["room_number"]]>1703) &(data[excel_column_name["room_number"]]!=666666)))&((data[excel_column_name["building_number"]]<6)&(data[excel_column_name["building_number"]]>36)))]
        
        for index,row in error_data.iterrows():
            building_number,room_number=(get_building_number_room_number_from_detial_adrees(row[excel_column_name["detail_address"]],if_upstream_park,if_use_pyqt,qtwidgets))
            building_number=int(building_number)
            room_number=int(room_number)
            if building_number>100 and building_number!=666666:
                messagebox_text= f'未能从"详细地址:""{row[excel_column_name["detail_address"]]}"中识别出{row[excel_column_name["custom_name"]]}订单的楼号，请检查后再试！'
                if if_use_pyqt:
                    QMessageBox.critical(qtwidgets,"错误对话框",messagebox_text)
                else:
                    print(messagebox_text)
                return False
            else:
                if building_number<6 and building_number>36:
                    if room_number<101 or room_number>1803:
                        messagebox_text= f'未能从"详细地址:""{row[excel_column_name["detail_address"]]}"中识别出{error_data[excel_column_name["custom_name"]].tolist()}订单的房号，请检查后再试！'
                        if if_use_pyqt:
                            QMessageBox.critical(qtwidgets,"错误对话框",messagebox_text)
                        else:
                            print(messagebox_text)
                        return False      
                       
            data.at[index,excel_column_name["building_number"]]=building_number
            data.at[index,excel_column_name["room_number"]]=room_number

            print(f'未能从"{excel_column_name["building_number"]}"或"{excel_column_name["room_number"]}"列中找到"{row[excel_column_name["custom_name"]]}"正确的楼号和室号，已从"{excel_column_name["detail_address"]}"列中重新解析出地址为{building_number}号{room_number}室')
            
        #修正别墅房号
        data_to_renew=data[(data[excel_column_name["building_number"]]>=6)&(data[excel_column_name["building_number"]]<=36)]
        for index,row in data_to_renew.iterrows():
            data.at[index,excel_column_name["room_number"]]=666666
            print(f'"{row[excel_column_name["custom_name"]]}"室号已更新为{row[excel_column_name["building_number"]]}号别墅。')
    #检查地址是否正确
    error_data=data[data[excel_column_name["building_number"]]==555555]
    if error_data.shape[0]>0:
        messagebox_text= f'未能识别{error_data[excel_column_name["custom_name"]].tolist()}订单的楼号，请检查后再试！'
        if if_use_pyqt:
            QMessageBox.critical(qtwidgets,"错误对话框",messagebox_text)
        else:
            print(messagebox_text)
        return False

    error_data=data[data[excel_column_name["room_number"]]==555555]
    if error_data.shape[0]>0:
        messagebox_text= f'未能识别{error_data[excel_column_name["custom_name"]].tolist()}订单的房号，请检查后再试！'
        if if_use_pyqt:
            QMessageBox.critical(qtwidgets,"错误对话框",messagebox_text)
        else:
            print(messagebox_text)
        return False

    print(f"共花费{(time.time()-start_time):0.1f}s完成订单预处理。")
    

    start_time2=time.time()
    # 输出派送单不带手机号
    reply=output_deliverynote_file(data, deliverynote_file_name, groupon_owner, product_name_list, excel_column_name,
                             max_row_number_per_page, page_margin_cm, show_sequence,title_sequence, if_hide_phone_number,if_upstream_park,if_use_pyqt,qtwidgets)
    if reply==True:
        print(f"共花费{(time.time()-start_time2):0.1f}s完成派送单的生成。")
    
    return reply
    